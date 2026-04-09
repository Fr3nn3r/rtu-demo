"""Generate professionally formatted case study docx for RTU Insurance Services."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# --- Colors ---
NAVY = RGBColor(0x0A, 0x1F, 0x3C)
DARK_BLUE = RGBColor(0x14, 0x3A, 0x6B)
MID_BLUE = RGBColor(0x1E, 0x5A, 0x96)
ACCENT = RGBColor(0x2B, 0x7A, 0xC6)
DARK_GRAY = RGBColor(0x2D, 0x2D, 0x2D)
MID_GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
QUOTE_BG = "EDF2F7"
SNAPSHOT_BG = "F0F5FA"
SNAPSHOT_BORDER = "1E5A96"
TABLE_HEADER_BG = "0A1F3C"
TABLE_ALT_BG = "F7F9FB"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_LIGHT = "D0D5DD"

doc = Document()

# --- Page setup (A4) ---
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# --- Default styles ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = DARK_GRAY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
style.paragraph_format.line_spacing = 1.25

for level, (sz, color, sp_before, sp_after) in enumerate([
    (Pt(22), NAVY, Pt(36), Pt(12)),
    (Pt(15), DARK_BLUE, Pt(24), Pt(8)),
    (Pt(12), MID_BLUE, Pt(18), Pt(6)),
], start=1):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.size = sz
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = sp_before
    h.paragraph_format.space_after = sp_after
    h.paragraph_format.keep_with_next = True


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tcPr.append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", BORDER_LIGHT)}"/>'
        )
        tcBorders.append(el)


def cell_text(cell, text, bold=False, color=None, size=Pt(10), align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color or DARK_GRAY


def add_quote(doc, text, attribution):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    bar = tbl.rows[0].cells[0]
    bar.width = Cm(0.3)
    set_cell_shading(bar, SNAPSHOT_BORDER)
    bar.text = ""
    bar.paragraphs[0].paragraph_format.space_before = Pt(0)
    bar.paragraphs[0].paragraph_format.space_after = Pt(0)
    content = tbl.rows[0].cells[1]
    content.width = Inches(5.5)
    set_cell_shading(content, QUOTE_BG)
    content.text = ""
    p = content.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run(f"\u201c{text}\u201d")
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = DARK_GRAY
    p2 = content.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(10)
    p2.paragraph_format.left_indent = Cm(0.4)
    run2 = p2.add_run(f"\u2014 {attribution}")
    run2.font.name = "Calibri"
    run2.font.size = Pt(9)
    run2.font.color.rgb = MID_GRAY
    run2.font.bold = True
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"val": "none", "sz": "0", "color": "FFFFFF"},
                bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                left={"val": "none", "sz": "0", "color": "FFFFFF"},
                right={"val": "none", "sz": "0", "color": "FFFFFF"})
    for row in tbl.rows:
        row.cells[0].width = Cm(0.3)
    doc.add_paragraph()


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 40)
    run.font.color.rgb = RGBColor(0xD0, 0xD5, 0xDD)
    run.font.size = Pt(6)


def add_hyperlink(paragraph, url, text, color="1E5A96", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}/>')
    new_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        f'    <w:color w:val="{color}"/><w:sz w:val="21"/>'
        f'    {"<w:u w:val=\"single\"/>" if underline else ""}'
        f'  </w:rPr>'
        f'  <w:t>{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_body(doc, text):
    doc.add_paragraph(text)


def add_body_bold_lead(doc, bold_part, rest):
    p = doc.add_paragraph()
    r = p.add_run(bold_part)
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = DARK_GRAY
    r2 = p.add_run(rest)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK_GRAY


# ============================================================
# HEADER / FOOTER
# ============================================================
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = hp.add_run("TrueAim.ai  |  Case Study")
run.font.size = Pt(8)
run.font.color.rgb = LIGHT_GRAY
run.font.name = "Calibri"

footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("trueaim.ai")
run.font.size = Pt(8)
run.font.color.rgb = LIGHT_GRAY
run.font.name = "Calibri"
fp.add_run("  |  ").font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
run = fp.add_run("Page ")
run.font.size = Pt(8)
run.font.color.rgb = LIGHT_GRAY
run.font.name = "Calibri"
run = fp.add_run()
run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
run = fp.add_run()
run._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
run = fp.add_run()
run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

# ============================================================
# TITLE
# ============================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("CASE STUDY")
run.font.name = "Calibri"
run.font.size = Pt(11)
run.font.color.rgb = ACCENT
run.font.bold = True
run.font.small_caps = True

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(20)
run = p.add_run("How RTU Insurance Services cut claim cycle time by 62% and stopped losing claims between systems")
run.font.name = "Calibri"
run.font.size = Pt(24)
run.font.color.rgb = NAVY
run.font.bold = True

# ============================================================
# SNAPSHOT BOX
# ============================================================
snapshot_data = [
    ("Company", "RTU Insurance Services (RTUSA)"),
    ("Industry", "Motor insurance underwriting management (P&C)"),
    ("Location", "South Africa"),
    ("Lines of business", "Taxi fleet motor claims: accident, theft, glass"),
    ("Insurer relationship", "Underwriting management agency for Renasa"),
    ("Challenge", "Claims tracked on Google Sheets across three disconnected systems. No workflow management, no SLA tracking."),
    ("Solution", "ClaimPilot by TrueAim \u2014 claims workflow platform with SLA monitoring, draft communications, and operational dashboards"),
    ("Results (12 months)", "62% shorter average cycle time, SLA compliance up from ~45% to 94%, assessor turnaround 3x faster, zero claims lost between systems"),
]

tbl = doc.add_table(rows=len(snapshot_data), cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(snapshot_data):
    row = tbl.rows[i]
    c0, c1 = row.cells[0], row.cells[1]
    c0.width = Cm(3.8)
    c1.width = Cm(10.0)
    set_cell_shading(c0, SNAPSHOT_BG)
    set_cell_shading(c1, SNAPSHOT_BG)
    cell_text(c0, label, bold=True, color=MID_BLUE, size=Pt(9.5))
    cell_text(c1, value, size=Pt(9.5))
    border_args = dict(
        left={"val": "single", "sz": "6", "color": SNAPSHOT_BORDER},
        right={"val": "single", "sz": "6", "color": SNAPSHOT_BORDER},
        bottom={"val": "single", "sz": "2", "color": "D6DFE8"},
    )
    if i == 0:
        border_args["top"] = {"val": "single", "sz": "6", "color": SNAPSHOT_BORDER}
    if i == len(snapshot_data) - 1:
        border_args["bottom"] = {"val": "single", "sz": "6", "color": SNAPSHOT_BORDER}
    set_cell_border(c0, **border_args)
    set_cell_border(c1, **border_args)

doc.add_paragraph()

# ============================================================
# SITUATION
# ============================================================
doc.add_heading("Situation", level=1)

add_body(doc,
    "RTU Insurance Services is an underwriting management agency in South Africa. They handle "
    "motor claims for taxi fleets insured through Renasa, one of the country\u2019s established "
    "short-term insurers. RTU\u2019s operations team manages every stage of the claim, from first "
    "notification through assessment, approval, repair or settlement, and salvage, all on "
    "Renasa\u2019s behalf."
)
add_body(doc,
    "At the time we engaged, RTU was growing quickly. Mike Mgodeli led sales and fleet strategy, "
    "Vassen Moodley oversaw operations, and Steve Cory handled the technical back-end. They had "
    "hundreds of active claims running at any given time and a team that moved fast."
)
add_body(doc, "Their systems had not kept up.")

# ============================================================
# CHALLENGE
# ============================================================
doc.add_heading("Challenge", level=1)

add_body(doc,
    "RTU\u2019s process ran on three systems that were never designed to talk to each other:"
)

doc.add_paragraph("Zoho, the CRM, where claim forms came in via web submission and phone", style="List Bullet")
doc.add_paragraph("Nimbus, Renasa\u2019s policy admin system, for checking policy validity and premium payments", style="List Bullet")
doc.add_paragraph("Rock, Renasa\u2019s claims admin system, for registering claims, appointing assessors, and processing payments", style="List Bullet")

add_body(doc,
    "None of these did workflow management. None tracked SLAs. None told an operator what to do next."
)
add_body(doc,
    "So RTU tracked everything on a Google Sheet. Every claim required an operator to leave one "
    "system, log into another, copy data between screens, then come back and update the spreadsheet. "
    "Milestones like assessor appointments, report deadlines, and QA referrals were tracked by memory."
)

add_quote(doc,
    "We had no tool that tells us \u2018this claim has been sitting in this status for too long\u2019 "
    "or \u2018here is your next best action.\u2019 The insurer\u2019s systems just don\u2019t do that "
    "from an underwriting manager\u2019s perspective. We were flying blind on SLAs.",
    "Steve Cory, Technical Lead, RTU Insurance Services"
)

add_body(doc, "Five problems kept surfacing:")

add_body_bold_lead(doc, "Assessor follow-up was reactive. ",
    "The 12-hour appointment SLA and 48-hour report SLA for assessors were unenforceable. RTU had "
    "no way to know which assessors were approaching a deadline, let alone which had already "
    "breached. Follow-up depended on someone remembering to check."
)
add_body_bold_lead(doc, "Claims disappeared between systems. ",
    "A claim registered on Nimbus might not make it to Rock for days. An assessor report received "
    "by email might sit in an inbox without anyone acting on it. With data entry spread across "
    "Zoho, Nimbus, Rock, and the Google Sheet, things fell through the cracks regularly."
)
add_body_bold_lead(doc, "Management had no visibility. ",
    "Vassen had no dashboard, no metrics, and no way to see which claims were at risk without "
    "reviewing the spreadsheet line by line."
)
add_body_bold_lead(doc, "Customer communication was inconsistent. ",
    "Policyholders and brokers received no proactive updates. When a claim fell within excess, the "
    "notification might arrive days late or not at all. Rejections had no standard communication process."
)
add_body_bold_lead(doc, "Theft investigations were hard to track. ",
    "Investigator reports carry a 14-day SLA. Without system tracking, these claims often sat for weeks."
)

add_body(doc,
    "RTU estimated their operators spent 40-50% of their working time on administrative overhead: "
    "copying data between systems, updating the spreadsheet, drafting emails by hand, and chasing "
    "assessors by phone. That was time not spent on adjudication, customer service, or anything "
    "that actually moved claims forward."
)
add_body(doc,
    "Missed SLAs were damaging the relationship with Renasa. Slow cycle times frustrated "
    "policyholders. And as the book grew, the manual process was becoming a ceiling on how much "
    "business RTU could handle."
)

# ============================================================
# ACTION
# ============================================================
doc.add_heading("Action", level=1)
doc.add_heading("Discovery and rapid prototyping (week 1\u20132)", level=2)

add_body(doc,
    "We ran a structured workshop with RTU\u2019s full leadership team. The goal was to map the "
    "actual workflow directly from the people who did it every day: every system handoff, every "
    "manual step, every decision point."
)
add_body(doc,
    "Mike was clear about what he wanted: \u201cIf we find a solution with you, we intend to run "
    "all our claims in your workflow and manage the entire claims process end-to-end within your "
    "space.\u201d He was equally clear about what was realistic: \u201cFor Version 1, assume no "
    "immediate integration with Rock or Nimbus. These are legacy systems with slow integration "
    "funnels. We need a tool that prompts humans to perform manual actions and tracks the "
    "timeline.\u201d"
)
add_body(doc,
    "That shaped the whole approach. Instead of waiting months for API integrations that might "
    "never arrive, we designed ClaimPilot around manual bridge steps. These are structured prompts "
    "that walk the operator through each interaction with Nimbus or Rock, give them copy buttons "
    "for the data they need, and track how long each step takes. When API integrations become "
    "available later, they slot in without changing the workflow."
)
add_body(doc,
    "Within two weeks we delivered a working interactive prototype covering all three claim types "
    "with their actual workflow differences:"
)

add_body_bold_lead(doc, "Accident claims ",
    "follow the full path: policy validation, Rock registration, assessor appointment, assessment "
    "received, auto-routing based on assessed amount versus excess (within excess, internal "
    "approval, or QA referral), then repair or total loss handling downstream."
)
add_body_bold_lead(doc, "Theft claims ",
    "are different. They skip excess collection, amount-based routing, and internal approval "
    "entirely. Every theft claim goes straight to mandatory QA because of the higher inherent risk."
)
add_body_bold_lead(doc, "Glass claims ",
    "are short: validation, registration, glass repairer appointment, repair done."
)

doc.add_heading("What ClaimPilot does", level=2)

add_body(doc,
    "ClaimPilot is a multi-tenant SaaS platform for claims workflow management. RTU was the "
    "first tenant. All the configuration (SLA thresholds, approval limits, claim types, workflow "
    "steps) is scoped to their operations."
)
add_body(doc, "Seven things it gave them:")

add_body_bold_lead(doc, "Document extraction. ",
    "Upload a claim form and the extraction pipeline reads it, populating up to 61 fields for "
    "theft forms, 54 for accident, 21 for glass. The operator reviews and corrects instead of "
    "typing from scratch."
)
add_body_bold_lead(doc, "Workflow management. ",
    "Every claim moves through a defined sequence of states with clear actions at each step. The "
    "system shows what needs to happen next, who is responsible, and logs every transition in an "
    "audit trail."
)
add_body_bold_lead(doc, "SLA monitoring. ",
    "Each step has a configurable SLA in calendar hours. The engine watches all active claims and "
    "flags them green (under 75% elapsed), amber (75-99%), or red (breached)."
)
add_body_bold_lead(doc, "Draft communications. ",
    "At 12 workflow milestones, ClaimPilot generates email drafts pre-filled with claim data: "
    "assessor appointments, SLA warnings, excess notifications, rejection notices, repair updates, "
    "settlement confirmations. The operator reviews, copies to their email client, and marks as sent."
)
add_body_bold_lead(doc, "Assessor and investigator tracking. ",
    "The system sends escalating reminders as deadlines approach (75%, 90%, 100%, then recurring "
    "after breach). No more relying on someone\u2019s memory."
)
add_body_bold_lead(doc, "Operational dashboard. ",
    "Claims by status, SLA compliance by step, assessor turnaround times, operator workload, and "
    "average time-to-close by claim type. All live."
)
add_body_bold_lead(doc, "Prioritised work queue. ",
    "The \u201cMy Queue\u201d view sorts an operator\u2019s claims by urgency: breached first, then "
    "approaching, then within SLA. The most urgent claim is always at the top."
)

doc.add_heading("Go-live (month 1\u20132)", level=2)

add_body(doc,
    "RTU used a clean cut-over: new claims went into ClaimPilot from day one, existing claims "
    "stayed on the Google Sheet until they closed out naturally. No migration, no risk."
)
add_body(doc,
    "In the first week, operators noticed claim creation time dropping. What used to take 15-20 "
    "minutes of manual data entry across systems was now 3-5 minutes of reviewing extracted data "
    "and confirming."
)
add_body(doc,
    "By the end of month two, nobody was updating the Google Sheet. Every active claim lived in "
    "ClaimPilot."
)

# ============================================================
# RESULTS
# ============================================================
doc.add_heading("Results", level=1)
doc.add_heading("12-month numbers", level=2)

results_headers = ["Metric", "Before", "After 12 months", "Change"]
results_data = [
    ("Average cycle time, accident", "34 days", "13 days", "-62%"),
    ("Average cycle time, theft", "48 days", "21 days", "-56%"),
    ("Average cycle time, glass", "8 days", "2.5 days", "-69%"),
    ("SLA compliance, all steps", "~45% (est.)", "94%", "+49 pp"),
    ("Assessor report turnaround", "5.2 days", "1.8 days", "-65%"),
    ("Investigator report turnaround", "22 days", "12 days", "-45%"),
    ("Operator time on admin", "40-50%", "~15%", "-30 pp"),
    ("Claims lost between systems", "3-5/month", "0", "Gone"),
    ("Policyholder comms within SLA", "Unmeasured", "97%", "Now tracked"),
    ("Active claims per operator", "~60", "~120", "2x"),
]

tbl = doc.add_table(rows=1 + len(results_data), cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
col_widths = [Cm(5.5), Cm(2.8), Cm(2.8), Cm(2.8)]

for i, h in enumerate(results_headers):
    cell = tbl.rows[0].cells[i]
    cell.width = col_widths[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(h)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.color.rgb = WHITE
    set_cell_shading(cell, TABLE_HEADER_BG)

for r_idx, (metric, before, after, change) in enumerate(results_data):
    row = tbl.rows[r_idx + 1]
    bg = TABLE_ALT_BG if r_idx % 2 == 0 else "FFFFFF"
    values = [metric, before, after, change]
    for c_idx, val in enumerate(values):
        cell = row.cells[c_idx]
        cell.width = col_widths[c_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(val)
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_GRAY
        if c_idx == 0:
            run.font.bold = True
        set_cell_shading(cell, bg)

doc.add_paragraph()

doc.add_heading("What actually changed", level=2)

add_body(doc,
    "SLA visibility had the biggest impact. Once management could see which claims were at risk "
    "and which assessors were underperforming, everything downstream improved. Vassen now opens "
    "the dashboard first thing every morning."
)
add_body(doc,
    "Assessors started self-correcting. Automated reminders at 75% and 90% of SLA meant most "
    "assessors responded before breach. The ones who consistently breached became visible in the "
    "data, so RTU could have real conversations backed by numbers and, where needed, shift work "
    "to faster assessors."
)
add_body(doc,
    "The team handled twice the volume without hiring. Once the admin overhead of copying data, "
    "tracking SLAs manually, and drafting emails from scratch went away, existing operators could "
    "carry 120 active claims instead of 60. The team went from struggling with 300 active claims "
    "to comfortably managing 600+."
)

add_quote(doc,
    "Before ClaimPilot, Monday mornings meant going through the spreadsheet line by line trying to "
    "figure out what fell through the cracks over the weekend. Now I open the dashboard and I can "
    "see exactly where every claim stands, which assessors need chasing, and which operators need "
    "support. We went from reactive to proactive overnight.",
    "Vassen Moodley, Operations Manager, RTU Insurance Services"
)

add_quote(doc,
    "We told TrueAim from day one, we want to run all our claims end-to-end in one system. Twelve "
    "months in, that\u2019s exactly what we\u2019re doing. The workflow management has changed how "
    "we operate. We\u2019re not just faster. We\u2019re visible. And that visibility is what lets "
    "us scale.",
    "Mike Mgodeli, Director, RTU Insurance Services"
)

doc.add_heading("Things RTU didn\u2019t expect", level=2)

add_body_bold_lead(doc, "Audit readiness. ",
    "The audit trail on every claim (status changes, field edits, document uploads, communications) "
    "gave RTU a complete timestamped record. When Renasa asked for claim handling evidence, RTU "
    "could produce it instantly instead of reconstructing timelines from emails."
)
add_body_bold_lead(doc, "Dispute resolution got faster. ",
    "When policyholders challenged a rejection or said they were never notified, RTU could point "
    "to the exact timestamp and the communication content. Disputes that used to take hours "
    "resolved in minutes."
)
add_body_bold_lead(doc, "Better reporting to Renasa. ",
    "With structured data on every claim, RTU started providing Renasa with performance reports "
    "that no other UMA in their network could match: SLA compliance rates, cycle times by claim "
    "type, assessor rankings. That made a difference in how Renasa viewed them as a partner."
)

doc.add_heading("Next phase", level=2)

add_body(doc, "RTU and TrueAim are now working on:")

doc.add_paragraph("Rock and Nimbus API integration to replace the manual bridge steps", style="List Bullet")
doc.add_paragraph("Fraud detection using TrueAim\u2019s document analysis pipeline", style="List Bullet")
doc.add_paragraph("Radics parts sourcing integration for automated price verification against repairer quotes", style="List Bullet")
doc.add_paragraph("Direct email sending to replace copy-to-clipboard", style="List Bullet")
doc.add_paragraph("Claims analytics built on 12 months of structured data: cost trending, leakage detection, predictive insights", style="List Bullet")

# ============================================================
# ABOUT TRUEAIM
# ============================================================
add_separator(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("About TrueAim")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.font.color.rgb = DARK_BLUE
run.font.bold = True

add_body(doc,
    "TrueAim AG is a Swiss insurtech company based in Zug, a spinout of Yarowa AG. We build "
    "workflow tools for P&C insurance claims teams: document extraction, workflow management, "
    "SLA enforcement, automated communications, and operational analytics."
)

p = doc.add_paragraph()
run = p.add_run("ClaimPilot is our claims workflow platform. To learn more, visit ")
run.font.name = "Calibri"
run.font.size = Pt(10.5)
run.font.color.rgb = DARK_GRAY
add_hyperlink(p, "https://trueaim.ai", "trueaim.ai")

# Disclaimer
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run(
    "This case study reflects projected outcomes based on the ClaimPilot MVP deployment for "
    "RTU Insurance Services. Metrics are modelled performance improvements derived from baseline "
    "operational data and system capabilities. Published March 2027."
)
run.font.name = "Calibri"
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = LIGHT_GRAY

# ============================================================
# SAVE
# ============================================================
import os
out = os.path.join(os.path.dirname(__file__), "case-study-rtusa.docx")
doc.save(out)
print(f"Saved: {out}")
